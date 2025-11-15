from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any
import io

# Almacenamiento en memoria para artefactos del informe
report_artifacts: Dict[str, Any] = {
    "plots": {},
    "summary": "No se han generado conclusiones todavía.",
    "eda_stats": None,
    "loaded_data_info": "No se ha cargado información de datos."
}

def clear_report_artifacts():
    """Limpia los artefactos para un nuevo informe."""
    report_artifacts["plots"].clear()
    report_artifacts["summary"] = "No se han generado conclusiones todavía."
    report_artifacts["eda_stats"] = None
    report_artifacts["loaded_data_info"] = "No se ha cargado información de datos."
    print("Artefactos de informe limpiados.")

def add_plot(name: str, plot_bytes: io.BytesIO):
    """Guarda un gráfico en memoria para el informe."""
    report_artifacts["plots"][name] = plot_bytes
    print(f"Gráfico '{name}' guardado para el informe.")

def set_summary(text: str):
    """Guarda el texto de resumen/conclusiones del agente."""
    report_artifacts["summary"] = text
    print("Resumen del informe guardado.")

def set_loaded_data_info(info: str):
    """Guarda información sobre los datos cargados."""
    report_artifacts["loaded_data_info"] = info

def set_eda_stats(stats: Dict[str, Any]):
    """Guarda las estadísticas del EDA."""
    report_artifacts["eda_stats"] = stats

def generate_report() -> io.BytesIO:
    """
    Genera un informe profesional de Word (.docx) con una estructura detallada.
    """
    document = Document()

    # --- Portada ---
    document.add_heading('Informe Analítico Profesional (IAIP)', 0)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Generado por el Sistema de Analítica de Datos Inteligente (SADI)').italic = True

    document.add_paragraph(f"Fecha de Generación: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC", style='Caption')
    document.add_page_break()

    # --- Tabla de Contenido (Placeholder) ---
    # La generación automática de ToC es compleja en python-docx.
    # Se añade un placeholder manual.
    document.add_heading('Tabla de Contenido', level=1)
    document.add_paragraph("1. Introducción y Objetivos")
    document.add_paragraph("2. Descripción de los Datos Cargados")
    document.add_paragraph("3. Resultados del Análisis Exploratorio (EDA)")
    document.add_paragraph("4. Resultados de los Modelos")
    document.add_paragraph("5. Conclusiones y Recomendaciones")
    document.add_page_break()

    # --- Sección 1: Introducción ---
    document.add_heading('1. Introducción y Objetivos', level=1)
    document.add_paragraph(
        "Este informe detalla el proceso y los hallazgos del análisis de datos realizado por el agente de IA. "
        "El objetivo es extraer insights, identificar patrones y construir modelos predictivos a partir de los datos proporcionados."
    )

    # --- Sección 2: Datos Cargados ---
    document.add_heading('2. Descripción de los Datos Cargados', level=1)
    document.add_paragraph(report_artifacts["loaded_data_info"])

    # --- Sección 3: EDA ---
    document.add_heading('3. Resultados del Análisis Exploratorio (EDA)', level=1)
    if report_artifacts["eda_stats"]:
        document.add_paragraph("A continuación se presentan las estadísticas descriptivas de las variables numéricas:")
        for col, stats in report_artifacts["eda_stats"].items():
            document.add_heading(f'Variable: {col}', level=2)
            table = document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Métrica'
            hdr_cells[1].text = 'Valor'
            for key, val in stats.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key.replace('_', ' ').title()
                row_cells[1].text = f'{val:.4f}' if isinstance(val, float) else str(val)
    else:
        document.add_paragraph("No se generaron estadísticas de EDA en este análisis.")

    # --- Sección 4: Modelos y Visualizaciones ---
    document.add_heading('4. Resultados de los Modelos y Visualizaciones', level=1)
    if not report_artifacts["plots"]:
        document.add_paragraph("No se generaron gráficos durante este análisis.")
    else:
        for name, plot_bytes in report_artifacts["plots"].items():
            try:
                document.add_heading(f'Gráfico: {name.replace("_", " ").title()}', level=2)
                plot_bytes.seek(0)
                document.add_picture(plot_bytes, width=Inches(6.0))
            except Exception as e:
                document.add_paragraph(f"No se pudo renderizar el gráfico '{name}'. Error: {e}")

    # --- Sección 5: Conclusiones ---
    document.add_heading('5. Conclusiones y Recomendaciones', level=1)
    document.add_paragraph(report_artifacts["summary"])

    # --- Guardar ---
    doc_buffer = io.BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)

    return doc_buffer
